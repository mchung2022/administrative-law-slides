import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

print("=== Generating 50+ Page Enterprise Master Tutorial Handbook for Antigravity AI ===")

doc = docx.Document()

# Page setup: A4, 1 inch margins
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

# Colors
NAVY = RGBColor(0x1B, 0x36, 0x5D)       # Primary Header (#1B365D)
SLATE = RGBColor(0x2C, 0x52, 0x82)      # Secondary Header (#2C5282)
AMBER = RGBColor(0xD9, 0x77, 0x06)      # Accent Color (#D97706)
TEXT_DARK = RGBColor(0x2D, 0x37, 0x48)  # Body Text (#2D3748)
MUTED_GREY = RGBColor(0x71, 0x80, 0x96)# Secondary text (#718096)

# Set Normal Style Font
style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Microsoft JhengHei'
font.size = Pt(11)
font.color.rgb = TEXT_DARK

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(15)
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(14)
    run.font.color.rgb = SLATE
    return p

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = NAVY
    return p

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = SLATE
    return p

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = AMBER
    return p

def add_p(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Microsoft JhengHei'
        r_pre.font.bold = True
        r_pre.font.color.rgb = NAVY
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_DARK
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Microsoft JhengHei'
        r_pre.font.bold = True
        r_pre.font.color.rgb = NAVY
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT_DARK
    return p

def add_callout(title, text, type_style='note'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.27)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    if type_style == 'warning':
        border_color = "D97706"
        bg_color = "FEF3C7"
        icon = "⚠️【重要警告與避坑指南】"
    elif type_style == 'tip':
        border_color = "2563EB"
        bg_color = "EFF6FF"
        icon = "💡【實務技巧與最佳做法】"
    elif type_style == 'prompt':
        border_color = "059669"
        bg_color = "ECFDF5"
        icon = "🤖【Antigravity 推薦提示詞範例 (Prompt Template)】"
    else:
        border_color = "1B365D"
        bg_color = "F0F4F8"
        icon = "📌【核心概念說明】"
        
    set_cell_background(cell, bg_color)
    
    # Left border styling
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="single" w:sz="36" w:space="0" w:color="{border_color}"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"{icon} {title}\n")
    run_t.font.name = 'Microsoft JhengHei'
    run_t.font.bold = True
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = NAVY
    
    run_b = p.add_run(text)
    run_b.font.name = 'Microsoft JhengHei'
    run_b.font.size = Pt(10)
    run_b.font.color.rgb = TEXT_DARK
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(4)

def add_code_block(code_text, lang="javascript"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.27)
    set_cell_margins(cell, top=100, bottom=100, left=180, right=180)
    set_cell_background(cell, "1E293B") # Dark Slate Code Bg
    
    # Border
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="6" w:space="0" w:color="334155"/><w:left w:val="single" w:sz="6" w:space="0" w:color="334155"/><w:bottom w:val="single" w:sz="6" w:space="0" w:color="334155"/><w:right w:val="single" w:sz="6" w:space="0" w:color="334155"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0) # Light slate
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(4)

print("[OK] Base docx helper functions initialized.")
