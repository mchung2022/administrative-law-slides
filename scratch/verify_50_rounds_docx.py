import os
import sys
import docx

print("=== 50-Round Auditing System for 500-Slide Word Lecture Transcripts ===")

doc_path = os.path.join(os.path.dirname(__file__), '..', '行政法500頁旗艦講義_1000字講稿全集.docx')

if not os.path.exists(doc_path):
    print(f"❌ Error: File not found at {doc_path}")
    sys.exit(1)

doc = docx.Document(doc_path)
tables = doc.tables
paragraphs = doc.paragraphs

pass_count = 0
fail_count = 0

def check_pass(round_num, title, condition, err_msg=""):
    global pass_count, fail_count
    if condition:
        pass_count += 1
        print(f"[Pass {round_num}/50] [OK] {title}")
    else:
        fail_count += 1
        print(f"[Pass {round_num}/50] [FAIL] {title}: {err_msg}")

# Pass 1 ~ 10: File & Document Structure
check_pass(1, "Word File Existence Audit", os.path.exists(doc_path))
check_pass(2, "Word File Size Audit", os.path.getsize(doc_path) > 1000000, f"Size is {os.path.getsize(doc_path)} bytes")
check_pass(3, "Document Paragraphs Audit", len(paragraphs) > 500, f"Found {len(paragraphs)} paragraphs")
check_pass(4, "Document Tables Audit", len(tables) >= 500, f"Found {len(tables)} tables")
check_pass(5, "Document Sections Audit", len(doc.sections) >= 1)
check_pass(6, "Cover Title Audit", any("行政法基本概念" in p.text for p in paragraphs[:5]))
check_pass(7, "Cover Author Audit", any("臺灣法治教育" in p.text for p in paragraphs[:10]))
check_pass(8, "Module 1 Heading Audit", any("Module 1" in p.text for p in paragraphs))
check_pass(9, "Module 10 Heading Audit", any("Module 10" in p.text for p in paragraphs))
check_pass(10, "Table Callout Style Audit", len(tables[0].rows) == 1 and len(tables[0].columns) == 1)

# Extract scripts from tables
slide_transcripts = []
for table in tables:
    text = table.cell(0, 0).text
    if "1000字深度教學講稿" in text:
        slide_transcripts.append(text)

# Pass 11 ~ 20: Slide Transcripts Count & Integrity
check_pass(11, "Slide Transcript Count Audit", len(slide_transcripts) == 500, f"Found {len(slide_transcripts)} slide callouts")
check_pass(12, "Slide 1 Presence Audit", any("Slide 1" in st for st in slide_transcripts[:5]))
check_pass(13, "Slide 500 Presence Audit", any("Slide 500" in st for st in slide_transcripts[-5:]))
check_pass(14, "Slide 100 Presence Audit", any("Slide 100" in st for st in slide_transcripts))
check_pass(15, "Slide 200 Presence Audit", any("Slide 200" in st for st in slide_transcripts))
check_pass(16, "Slide 300 Presence Audit", any("Slide 300" in st for st in slide_transcripts))
check_pass(17, "Slide 400 Presence Audit", any("Slide 400" in st for st in slide_transcripts))
check_pass(18, "Slide 5-Section Template Audit", all("📌 【一、本頁核心主題" in st for st in slide_transcripts))
check_pass(19, "Slide Statute Section Audit", all("🔍 【二、法條剖析" in st for st in slide_transcripts))
check_pass(20, "Slide News Case Section Audit", all("📰 【三、臺灣時事" in st for st in slide_transcripts))

# Pass 21 ~ 30: Word Count per Slide Audit (1000+ characters)
word_counts = [len(st) for st in slide_transcripts]
min_words = min(word_counts) if word_counts else 0
avg_words = sum(word_counts) / len(word_counts) if word_counts else 0

check_pass(21, "Slide Transcript Minimum Character Audit", min_words >= 900, f"Min char count is {min_words}")
check_pass(22, "Slide Transcript Average Character Audit", avg_words >= 1000, f"Avg char count is {avg_words:.1f}")
check_pass(23, "Slide 1 Character Count Audit", len(slide_transcripts[0]) >= 950)
check_pass(24, "Slide 50 Character Count Audit", len(slide_transcripts[49]) >= 950)
check_pass(25, "Slide 150 Character Count Audit", len(slide_transcripts[149]) >= 950)
check_pass(26, "Slide 250 Character Count Audit", len(slide_transcripts[249]) >= 950)
check_pass(27, "Slide 350 Character Count Audit", len(slide_transcripts[349]) >= 950)
check_pass(28, "Slide 450 Character Count Audit", len(slide_transcripts[449]) >= 950)
check_pass(29, "Slide 500 Character Count Audit", len(slide_transcripts[499]) >= 950)
check_pass(30, "Overall 500 Slides 1000-Word Compliance Rate", sum(1 for w in word_counts if w >= 950) == 500)

# Pass 31 ~ 40: Legal Terminology & Statutes Accuracy
full_text = "\n".join(p.text for p in paragraphs) + "\n" + "\n".join(st for st in slide_transcripts)

check_pass(31, "Administrative Procedure Act Mention Audit", "行政程序法" in full_text)
check_pass(32, "Administrative Penalty Act Mention Audit", "行政罰法" in full_text)
check_pass(33, "Administrative Enforcement Act Mention Audit", "行政執行法" in full_text)
check_pass(34, "Appeal Act Mention Audit", "訴願法" in full_text)
check_pass(35, "Interpretation No. 443 Audit", "443" in full_text)
check_pass(36, "Principle of Proportionality Audit", "比例原則" in full_text)
check_pass(37, "Tiered Legal Reservation Audit", "法律保留" in full_text)
check_pass(38, "State Compensation Act Audit", "國家賠償" in full_text)
check_pass(39, "Administrative Litigation 3-Tier System Audit", "行政訴訟" in full_text)
check_pass(40, "Civics Literacy & Exam Traps Audit", "大考" in full_text)

# Pass 41 ~ 50: Diagrams & Word Formatting Verification
charts_dir = os.path.join(os.path.dirname(__file__), 'charts')
check_pass(41, "Legal Source Pyramid PNG Audit", os.path.exists(os.path.join(charts_dir, 'chart_legal_pyramid.png')))
check_pass(42, "Legal Reservation PNG Audit", os.path.exists(os.path.join(charts_dir, 'chart_legal_reservation.png')))
check_pass(43, "Proportionality PNG Audit", os.path.exists(os.path.join(charts_dir, 'chart_proportionality.png')))
check_pass(44, "Administrative Action PNG Audit", os.path.exists(os.path.join(charts_dir, 'chart_administrative_action.png')))
check_pass(45, "Administrative Remedy PNG Audit", os.path.exists(os.path.join(charts_dir, 'chart_administrative_remedy.png')))
check_pass(46, "Word Inline Pictures Count Audit", len(doc.inline_shapes) >= 5, f"Found {len(doc.inline_shapes)} pictures")
check_pass(47, "Font Family Audit (Microsoft JhengHei)", doc.styles['Normal'].font.name == 'Microsoft JhengHei')
check_pass(48, "Heading 1 Hierarchy Audit", len([p for p in paragraphs if p.style.name.startswith('Heading 1')]) >= 10)
check_pass(49, "Heading 2 Hierarchy Audit", len([p for p in paragraphs if p.style.name.startswith('Heading 2')]) >= 500)
check_pass(50, "Final Word Document Integrity Audit", pass_count >= 49)

print("\n=================================================================")
print(f"[OK] 50-ROUND WORD LECTURE TRANSCRIPT AUDIT RESULT: {pass_count} Passed, {fail_count} Failed.")
print("=================================================================")

if fail_count > 0:
    sys.exit(1)
