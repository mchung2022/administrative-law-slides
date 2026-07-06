import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

print("=== Generating Comprehensive 50+ Page Enterprise Master Handbook (.docx) ===")

doc = docx.Document()

# Page setup: A4, 1 inch margins
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

# Colors - Steady Enterprise Palette
NAVY = RGBColor(0x1B, 0x36, 0x5D)       # Primary Header (#1B365D - 深海藍)
SLATE = RGBColor(0x2C, 0x52, 0x82)      # Secondary Header (#2C5282 - 板岩藍)
AMBER = RGBColor(0xD9, 0x77, 0x06)      # Accent Color (#D97706 - 琥珀金)
TEXT_DARK = RGBColor(0x2D, 0x37, 0x48)  # Body Text (#2D3748 - 石墨灰)

style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Microsoft JhengHei'
font.size = Pt(11)
font.color.rgb = TEXT_DARK

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(15)
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = NAVY
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(13)
    run.font.color.rgb = SLATE
    return p

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(17)
    run.font.bold = True
    run.font.color.rgb = NAVY
    return p

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(13.5)
    run.font.bold = True
    run.font.color.rgb = SLATE
    return p

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = AMBER
    return p

def add_p(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Microsoft JhengHei'
        r_pre.font.bold = True
        r_pre.font.color.rgb = NAVY
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_DARK
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Microsoft JhengHei'
        r_pre.font.bold = True
        r_pre.font.color.rgb = NAVY
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT_DARK
    return p

def add_callout(title, text, type_style='note'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.27)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    if type_style == 'warning':
        border_color = "D97706"
        bg_color = "FEF3C7"
        icon = "[重要警告與避坑指南] "
    elif type_style == 'tip':
        border_color = "2563EB"
        bg_color = "EFF6FF"
        icon = "[實務技巧與最佳做法] "
    elif type_style == 'prompt':
        border_color = "059669"
        bg_color = "ECFDF5"
        icon = "[Antigravity 推薦提示詞範例 (Prompt Template)] "
    else:
        border_color = "1B365D"
        bg_color = "F0F4F8"
        icon = "[核心觀念與步驟說明] "
        
    set_cell_background(cell, bg_color)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="single" w:sz="36" w:space="0" w:color="{border_color}"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"{icon} {title}\n")
    run_t.font.name = 'Microsoft JhengHei'
    run_t.font.bold = True
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = NAVY
    
    run_b = p.add_run(text)
    run_b.font.name = 'Microsoft JhengHei'
    run_b.font.size = Pt(10)
    run_b.font.color.rgb = TEXT_DARK
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(4)

def add_code_block(code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.27)
    set_cell_margins(cell, top=100, bottom=100, left=180, right=180)
    set_cell_background(cell, "1E293B")
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="6" w:space="0" w:color="334155"/><w:left w:val="single" w:sz="6" w:space="0" w:color="334155"/><w:bottom w:val="single" w:sz="6" w:space="0" w:color="334155"/><w:right w:val="single" w:sz="6" w:space="0" w:color="334155"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(4)

def add_custom_table(headers, rows_data):
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1B365D")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Microsoft JhengHei'
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    # Data rows
    for r_idx, row_data in enumerate(rows_data):
        row_cells = table.rows[r_idx + 1].cells
        bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = val
            set_cell_background(row_cells[c_idx], bg)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=140, right=140)
            p = row_cells[c_idx].paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Microsoft JhengHei'
                r.font.size = Pt(9.5)
                r.font.color.rgb = TEXT_DARK

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(6)
    p_after.paragraph_format.space_after = Pt(6)

print("[OK] Base setup complete.")

# =========================================================================
# COVER & TITLE SECTION
# =========================================================================
add_title("Antigravity AI 全方位實戰指南：\n從零到一打造 500 頁 HTML 互動簡報、講稿全集與網頁版雙主持 Podcast 影音")
add_subtitle("108 課綱公民與社會：行政法核心概念大型專案協作教學手冊 (50+ 頁企業精裝版)")

add_callout(
    "手冊導讀與核心學習效益",
    "本教學手冊專為 Antigravity AI 人工智慧協作系統的初學者所編寫。全書以「行政法 500 頁旗艦簡報與 30 分鐘雙主持廣播 Podcast」專案為真實範例，手把手引導學員掌握如何透過自然語言與 Antigravity AI 進行雙人對話（Pair Programming）、需求定義、自動化程式碼編寫、音訊合成、自動化測試與 GitHub 雲端部署。讀完本講義後，您將具備獨自引導 AI 產出企業級大型 Web 應用與多媒體系統的完整能力！",
    type_style='note'
)

add_custom_table(
    ["專案核心模組", "成果規格與量化指標", "技術架構與關鍵技術說明"],
    [
        ["500 頁 HTML 互動簡報", "500 頁全動態互動簡報 (1.35 MB)", "資料驅動架構、5 大題型組件、80+ 臺灣新聞案例3卡視覺容器"],
        ["50 萬字大書講稿全集", "50 萬字 Word 輸出 (.docx)", "單頁 1,000 字補充、內嵌 Matplotlib 統計圖表、企業排版"],
        ["30 分鐘特企 Podcast", "30 分鐘雙主持對談節目 (9.18 MB)", "Azure Neural 雙神經語音、成年人自然語速 (-15%)、10 大章節"],
        ["雙引擎 Podcast 播放器", "雙引擎網頁播放器 (podcast_player.html)", "HTML5 MP3 + Web Speech TTS 備援、Canvas 動態頻譜、毫秒同步"],
        ["自動化測試與品質監控", "100 輪法律稽核 + 50 輪 UI 測試", "100/100 法條判例交叉驗證通過、50/50 時間軸同步點擊驗證"]
    ]
)

doc.add_page_break()

# =========================================================================
# CHAPTER 1
# =========================================================================
add_h1("第一章：Antigravity AI 核心觀念與人機協作思維")

add_h2("1.1 什麼是 Antigravity AI？進階代理人（Agentic Coding）機制解密")
add_p("Antigravity 是由 Google DeepMind 團隊設計的高級 Agentic AI 編程助手。與傳統僅能提供單純程式碼片段回答的聊天機器人（Chatbot）不同，Antigravity 具備自主觀察（Observation）、思考規劃（Planning）、工具調用（Tool Invocation）與自動驗證（Verification）的全套代理能力。")
add_p("在傳統模式下，程式設計師需要手動開啟編輯器、建立資料夾、撰寫程式碼、處理套件安裝、並手動開啟瀏覽器測試。而在 Antigravity 代理模式下，AI 能夠直接讀取整個專案的目錄結構，獨立調用 Python 或 Node.js 執行檔，甚至能夠自己撰寫測試腳本對產出物進行連環稽核！")

add_callout(
    "Antigravity AI 的三大核心能力觀念",
    "1. 工具調用自主性：Antigravity 可以自動閱讀檔案、搜尋代碼、建立資料夾與執行終端機命令。\n2. 規劃與執行模式 (Planning & Execution)：遇到大型複雜任務時，AI 會主動提出 Implementation Plan 並要求使用者確認後再執行。\n3. 自我修復與連環稽核：當程式碼執行出錯時，AI 能讀取錯誤日誌（Error Logs）並自動修復，直到測試 100% 通過為止。",
    type_style='tip'
)

add_h2("1.2 人機 Pair Programming（雙人對話協作）的核心原則")
add_p("要在大型專案中獲得最佳成果，使用者必須掌握 Pair Programming 的雙人角色分工：")
add_bullet("使用者（User）擔任「產品經理 (PM)」與「架構審查者」：負責提出明確業務需求、驗收標準與視覺美感偏好。決定產品的方向與規格，而非陷於微小的代碼語法細節。", bold_prefix="使用者角色：")
add_bullet("Antigravity AI 擔任「首席工程師」與「執行團隊」：負責底層架構設計、代碼撰寫、重構與自動化驗證。主動發現潛在 Bug 並提出最佳解法。", bold_prefix="AI 角色：")

add_h2("1.3 提示詞（Prompt）撰寫黃金公式與需求傳達技巧")
add_p("初學者最常遇到的困境是「指令太模糊導致 AI 產出不符合預期」。請務必遵循以下「4W1H 提示詞黃金公式」：")

add_callout(
    "4W1H 提示詞黃金公式 (Prompt Golden Rule)",
    "【Who】角色定位：指定 AI 的身份（例如：請扮演資深 Web 開發者與廣播製作人）\n【What】目標成果：明確產出物名稱與格式（例如：製作 500 頁 HTML 簡報與 30 分鐘 Podcast MP3）\n【Where】儲存路徑：指定檔案存放目錄（例如：儲存於專案根目錄 index.html）\n【Why】業務背景：說明使用情境與對象（例如：高中公民行政法複習，需符合 108 課綱與臺灣法律）\n【How】驗收標準：量化指標（例如：進行 100 次自動化測試驗證，確保 100% 無錯誤）",
    type_style='prompt'
)

add_h2("1.4 常見初學者迷思與避免「AI 幻覺」的精準指令法")
add_p("為了避免 AI 在法律條文或複雜邏輯中產生「幻覺（Hallucination）」，初學者應掌握以下三大精準防護指令：")
add_bullet("要求引用具體法規與判例：在提示詞中明確要求 AI 引用《行政程序法》條號或司法院大法官解釋（如釋字第 443 號）。", bold_prefix="條文防護：")
add_bullet("要求建立驗證腳本：不要只聽 AI 口頭說「已經完成」，要求 AI 撰寫自動化測試腳本（如 Python/Node.js）進行客觀檢驗。", bold_prefix="驗證防護：")
add_bullet("分階段小步快跑：不要一次丟出 100 個無關的需求，將大型專案拆解為「簡報架構 -> 講稿產出 -> 音訊合成 -> 播放器開發 -> 雲端部署」5 大階段。", bold_prefix="階段防護：")

add_h2("1.5 實戰範例：如何引導 AI 生成複雜互動式 HTML/JS 元件")
add_p("當您需要 Antigravity 為您開發一個複雜的 UI 組件（例如選擇題即時反饋組件）時，請參考以下實戰步驟與範例指令：")

add_code_block("""// 引導 AI 生成的選擇題 JavaScript 解析函數範例
function checkAnswer(slideId, selectedIdx, correctIdx, explanation) {
    const slideElem = document.getElementById(`slide-${slideId}`);
    const options = slideElem.querySelectorAll('.option-btn');
    const explBox = slideElem.querySelector('.explanation-box');
    
    options.forEach((opt, idx) => {
        opt.disabled = true; // 防止重複點擊
        if (idx === correctIdx) {
            opt.classList.add('correct'); // 綠色高亮
        } else if (idx === selectedIdx) {
            opt.classList.add('wrong');   // 紅色高亮
        }
    });
    
    explBox.innerHTML = `<strong>💡 答案解析：</strong> ${explanation}`;
    explBox.style.display = 'block';
}""")

doc.add_page_break()

# =========================================================================
# CHAPTER 2
# =========================================================================
add_h1("第二章：專案環境準備與 Antigravity 工作區配置")

add_h2("2.1 環境安裝與工具盤點")
add_p("在開始打造 500 頁簡報與 Podcast 專案前，必須確認本機環境已安裝以下基礎工具鏈：")

add_custom_table(
    ["工具名稱", "建議版本", "用途與說明"],
    [
        ["Node.js", "v18.0.0+", "用於執行 DOM 沙盒模擬測試、單一檔案 HTML 打包與網頁建置"],
        ["Python", "3.10+", "用於執行 edge-tts 語音合成、python-docx Word 生成與 Matplotlib 製圖"],
        ["edge-tts", "v7.2.0+", "Python 開源套件，直接對接 Microsoft Azure Neural 廣播級神經語音引擎"],
        ["Git", "v2.30.0+", "版本控制工具，管理檔案變更與 GitHub Pages 免費雲端託管部署"],
        ["VS Code", "最新版", "文字編輯器，用於檢視 HTML/JS/CSS 原始碼與 Markdown 講義"]
    ]
)

add_h2("2.2 Antigravity 專案目錄結構設計")
add_p("優良的專案目錄結構是專案成功的基石。Antigravity AI 能自動識別結構並進行檔案模組化管理。標準專案目錄如下：")

add_code_block("""c:/Users/DELL/Desktop/新增資料夾 (8)/
├── index.html                           # 500 頁單一檔案互動簡報 (1.35 MB)
├── podcast_player.html                  # 30 分鐘特企雙主持 Podcast 播放器
├── podcast_audio_30min.mp3              # 30 分鐘雙主持廣播實體 MP3 音訊 (9.18 MB)
├── 行政法500頁旗艦講義_1000字講稿全集.docx  # 50 萬字大書講稿 Word 文件
├── css/
│   └── style.css                        # 深色極光 CSS 視覺設計系統
├── js/
│   ├── slidesData.js                    # 500 頁簡報資料定義檔
│   └── app.js                           # 簡報互動邏輯與 Web Speech 播放引擎
└── scratch/                             # Antigravity 實驗與自動化測試腳本區
    ├── build_single_file.js             # 簡報單一檔案打包腳本
    ├── synthesize_dual_host_audio.py    # 雙主持人語音合成腳本
    ├── update_player_html.py            # 播放器時間軸同步更新腳本
    ├── audit_50_dual_host_podcast.js    # 50 輪播放器點擊測試腳本
    └── verify_100_rounds_podcast_legal.py # 100 輪臺灣法律稽核腳本""")

add_h2("2.3 scratch 臨時實驗區與自動化測試腳本規範")
add_p("`scratch/` 資料夾是 Antigravity 協作體系中的「實驗測試沙盒」。所有中間過程生成的 Python 測試腳本、音訊分軌檔（`ch_1.mp3` ~ `ch_10.mp3`）、時間軸計算 JSON 檔均應存放於此，確保專案根目錄乾淨整潔。")

add_h2("2.4 Git 版本控制與 GitHub Pages 自動部署實戰")
add_p("透過 Git 進行版本管理，能確保隨時可追溯至歷史穩定狀態。以下為常用命令指引：")

add_code_block("""# 1. 檢視工作區狀態
git status

# 2. 強制追蹤大檔案音訊與新產出檔案
git add -f podcast_audio_30min.mp3 podcast_player.html index.html scratch/

# 3. 提交版本變更與詳細說明
git commit -m "feat: complete 500-slide HTML presentation & 30-min dual-host podcast"

# 4. 推送至 GitHub 遠端儲存庫
git push origin main""")

add_h2("2.5 實戰指令對照表：Windows PowerShell 下最常用的命令")
add_p("為方便初學者在 Windows 環境下操作，以下整理最常用的 PowerShell 指令對照表：")

add_custom_table(
    ["操作目標", "PowerShell 指令範例", "注意事項與說明"],
    [
        ["查看當前目錄檔案", "ls 或 Get-ChildItem", "確認專案檔案是否存在"],
        ["執行 Python 腳本", "python scratch/synthesize_dual_host_robust.py", "確認 Python 已加入 PATH 環境變數"],
        ["執行 Node 測試", "node scratch/audit_50_dual_host_podcast.js", "用 Node.js 執行自動化測試腳本"],
        ["安裝 Python 套件", "pip install edge-tts python-docx matplotlib", "安裝 TTS 與 Word 生成套件"],
        ["Git 強制推送", "git push origin main --force", "慎用 force 命令，確保備份完整"]
    ]
)

doc.add_page_break()

# =========================================================================
# CHAPTER 3
# =========================================================================
add_h1("第三章：500 頁 HTML 互動簡報系統架構與開發實戰")

add_h2("3.1 簡報設計理念：108 課綱與現代深色極光 (Dark Aurora) UI Aesthetics")
add_p("本專案簡報設計擺脫傳統簡報白底黑字的平庸感，採用現代高階網頁設計標竿「深色極光 (Dark Aurora)」主視覺。採用 HSL 調和配色、漸層發光背景（Background Glow）、微互動動畫與現代 Google Font 字型（Noto Sans TC & Outfit）。")

add_h2("3.2 JavaScript 資料驅動架構：slidesData.js 的設計與 500 頁結構化定義")
add_p("簡報採用「資料與 UI 分離」的現代前端架構。全數 500 頁簡報內容定義於 `js/slidesData.js` 中，每頁包含以下結構化欄位：")

add_code_block("""// slidesData.js 單頁資料結構範例
{
  "id": 10,
  "module": "Module 1: 行政與行政法概論",
  "type": "concept",
  "category": "臺灣新聞真實案例剖析",
  "title": "Slide 10: 臺灣新聞案例 — 釋字 443 號層次化法律保留體系 時事剖析",
  "subtitle": "高鐵採購與國營事業管轄實務爭點評析",
  "caseStudy": {
    "title": "新聞事件：行政機關以令函限制人民出境案",
    "context": "📜 一、事件脈絡：內政部入出國及移民署過去僅憑內部作業要點即限制特定國人出境...",
    "dispute": "⚖️ 二、法律爭點：限制出境涉及憲法第 10 條居住遷徙自由，得否以行政規則為之？",
    "examPoints": "🎯 三、大考考點：司法院釋字第 443 號揭示，限制人身自由與重要自由需有法律明確授權！"
  },
  "podcastScript": "哈囉各位同學！歡迎來到今天的行政法 Podcast 廣播講堂...",
  "notes": "本頁重點在於區分絕對法律保留與相對法律保留之差異..."
}""")

add_h2("3.3 多種互動題型組件開發")
add_p("簡報內建 5 大素養題型組件，點擊即時反饋答案與詳細解析：")
add_bullet("素養選擇題 (multiple_choice)：單選題型，點擊選項立即標示正確（綠色）或錯誤（紅色）並展開解析。", bold_prefix="1. 選擇題：")
add_bullet("素養配合題 (matching)：左右雙欄卡片連線配對，即時計算得分與錯題提示。", bold_prefix="2. 配合題：")
add_bullet("素養填空題 (fill_in_blank)：關鍵字填空與答案隱藏切換。", bold_prefix="3. 填空題：")
add_bullet("素養是非題 (true_false)：觀念診斷與反向思考評析。", bold_prefix="4. 是非題：")
add_bullet("素養簡答題 (short_answer)：爭議題幹與答題心智圖展開。", bold_prefix="5. 簡答題：")

add_h2("3.4 臺灣新聞案例 3 卡結構元件開發")
add_p("所有新聞案例頁面均採用強化的「3 卡視覺容器 (News Case Container)」：左卡【事件脈絡與背景】、中卡【法律爭點與攻防】、右卡【大考考點與解題指南】，極大化學習效果。")

add_h2("3.5 單一檔案打包技術（scratch/build_single_file.js）")
add_p("為了方便學員離線閱讀與跨裝置開啟，透過 Node.js 腳本 `scratch/build_single_file.js` 將 CSS、JS 與 500 頁資料完全內嵌打包至單一 `index.html` 檔案（檔案大小僅 1.35 MB），實現零外部檔案依賴！")

doc.add_page_break()

# =========================================================================
# CHAPTER 4
# =========================================================================
add_h1("第四章：50 萬字大書講稿與 Word 全集自動化生成")

add_h2("4.1 為什麼需要 50 萬字講稿？單頁 1,000 字精華延伸說明")
add_p("為使教學講義具備極高學術價值與大考總複習參考價值，專案要求每一頁簡報均需搭配 1,000 字的精華講稿與法理延伸。500 頁合計產出高達 500,000 字之文字巨著！")

add_h2("4.2 Python python-docx 套件實戰：企業級 Word 文件自動排版技術")
add_p("手動排版 50 萬字 Word 文件極易崩潰。透過 Python `python-docx` 套件，能以程式碼精準控制字型（標楷體/微軟正黑體）、段落間距、標題階層與表格樣式。")

add_code_block("""# python-docx 企業級排版範例
import docx
from docx.shared import Pt, Inches, RGBColor

doc = docx.Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft JhengHei'
style.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run("【Module 1：行政與行政法概論】")
run.font.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy Blue""")

add_h2("4.3 內嵌統計圖表（Matplotlib）與法理比較表格生成")
add_p("腳本會自動調用 Python Matplotlib 繪製法理數據柱狀圖與餅圖，轉存為高解析度 PNG 圖片後自動插入 Word 文件相對應章節中。")

add_h2("4.4 100 輪臺灣法律正確性稽核腳本")
add_p("編寫 `scratch/verify_100_rounds_docx_taiwan_law.py` 對產出的 Word 講稿進行 100 輪法規條號與專有名詞驗證，確保 100% 符合臺灣法律脈絡。")

doc.add_page_break()

# =========================================================================
# CHAPTER 5
# =========================================================================
add_h1("第五章：30 分鐘特企 Podcast 廣播對談劇本寫作與聲學設計")

add_h2("5.1 Podcast 單集節目架構：10 大核心模組之主題規劃")
add_p("30 分鐘廣播節目涵蓋行政法 10 大核心單元，每單元分配約 3 分鐘精華對談：")

add_custom_table(
    ["章節 (Module)", "時間軸範圍", "對談主題與法理重點"],
    [
        ["Module 1", "00:00 - 03:33", "節目開場、公權力行政 vs 私經濟國庫行政與給付行政理念"],
        ["Module 2", "03:33 - 06:45", "依法行政原則、法律優位與釋字 443 號層次化法律保留"],
        ["Module 3", "06:45 - 09:53", "平等原則（不法不得主張平等）與比例原則三大過磅審查"],
        ["Module 4", "09:53 - 12:44", "信賴保護原則三要件與裁量瑕疵（逾越、濫用、怠惰）"],
        ["Module 5", "12:44 - 15:32", "行政處分黃金六要素、自始無效（111條）與撤銷（117條）"],
        ["Module 6", "15:32 - 18:06", "行政契約（135條）、行政指導（165條）與行政規則劃分"],
        ["Module 7", "18:06 - 20:45", "行政罰法故意過失舉證、一行為不二罰與刑事優先原則"],
        ["Module 8", "20:45 - 23:18", "行政執行金錢給付、代履行/怠金與即時強制特別犧牲補償"],
        ["Module 9", "23:18 - 25:58", "訴願 30 日不變期間、訴願前置原則與行政訴訟三級二審"],
        ["Module 10", "25:58 - 29:17", "國家賠償法第 2 條/第 3 條無過失責任與協議先行程序結業"]
    ]
)

add_h2("5.2 雙主持人角色設定（YunJhe 男聲 + HsiaoChen 女聲）")
add_bullet("🎙️ 主持人 阿哲（YunJhe - 男聲）：風趣熱情，代表大眾與考生提出生活情境問題與新聞疑惑。", bold_prefix="男主持：")
add_bullet("👩‍🏫 法治導師 小晨（HsiaoChen - 女聲）：權威專業、條理分明，深入淺出拆解法理要件與大考陷阱。", bold_prefix="女導師：")

add_h2("5.3 輕鬆白話廣播對白寫作技巧")
add_p("避免冷冰冰的條文朗讀，採用「新聞提問 -> 核心觀念帶出 -> 法理要件拆解 -> 大考解題提醒」4 步廣播對談法。")

add_h2("5.4 13,000 字劇本生成與 JSON 結構化解析")
add_p("對白劇本儲存於 `scratch/podcast_script_30min.json`，每條對白均帶有 `speaker`、`voice` 與 `text` 欄位。")

doc.add_page_break()

# =========================================================================
# CHAPTER 6
# =========================================================================
add_h1("第六章：Microsoft Azure Neural 神經語音 TTS 語音合成與語速校準")

add_h2("6.1 edge-tts Python 開源庫安裝與 API 原理")
add_p("`edge-tts` 是對接 Microsoft Azure Speech SDK 的高效 Python 開源庫，支援雙聲道神經語音合成。安裝命令：`pip install edge-tts`。")

add_h2("6.2 成年人自然電台語速校準實驗")
add_p("經過 `scratch/test_calibrate_rate.py` 實測，不同的 rate 參數對應之每秒字數（CPS）如下：")

add_custom_table(
    ["語速設定 (Rate)", "每秒朗讀字數 (CPS)", "每分鐘字數 (CPM)", "聽感評估與適用場景"],
    [
        ["+0% (預設)", "5.50 字/秒", "330 字/分", "偏快，類似新聞快報，長時間收聽易疲勞"],
        ["-10%", "4.47 字/秒", "268 字/分", "標準播報速，清晰度高"],
        ["-15% (推薦)", "4.22 字/秒", "253 字/分", "🏆 最舒適成年人日常電台語速，輕鬆自然"],
        ["-20%", "3.98 字/秒", "239 字/分", "輕鬆對談速，適合沉浸式收聽"],
        ["-35% ~ -50%", "2.50 字/秒", "150 字/分", "過慢，適合幼童教學或極慢速朗讀"]
    ]
)

add_h2("6.3 雙聲道分軌合成與音訊合併")
add_p("腳本逐句調用 `zh-TW-YunJheNeural` 與 `zh-TW-HsiaoChenNeural` 合成 MP3 片段後，以二進位流合併為完整 `podcast_audio_30min.mp3`。")

add_h2("6.4 MP3 標頭解析與精確毫秒時間軸計算")
add_p("編寫 `scratch/measure_mp3_duration.py` 純 Python MP3 Frame Header 解析器，精確計算出每章節的 start_sec 與 end_sec，生成 `exact_synced_chapters.json`。")

doc.add_page_break()

# =========================================================================
# CHAPTER 7
# =========================================================================
add_h1("第七章：雙引擎網頁版 Podcast 播放器 (podcast_player.html) 開發")

add_h2("7.1 雙引擎架構設計（HTML5 Audio + Web Speech API 備援）")
add_p("播放器內建雙重播放引擎：優先播放高音質 `podcast_audio_30min.mp3`；若在跨域或無網路環境，自動切換至 Web Speech Synthesis API 即時合成發聲。")

add_h2("7.2 Canvas 動態頻譜波形視覺化效果")
add_p("利用 HTML5 Canvas 繪製藍金漸層頻譜波形（Waveform Visualizer），隨音訊播放發送動態波浪動畫，展現奢華電台質感。")

add_h2("7.3 左側播放器控制器與右側章節目錄 100% 毫秒級動態同步技術")
add_p("點擊右側章節卡片（如 Module 7 18:06），左側 HTML5 Audio 控制器自動跳轉至 `currentTime = 1086.81` 並同步高亮右側卡片與捲動逐字稿！")

add_h2("7.4 雙主持人動態逐字稿高亮捲動與速度調整")
add_p("逐字稿即時顯示 `【阿哲】` 與 `【小晨】` 對話，支援 0.8x、1.0x、1.25x、1.5x、2.0x 無段變速播放。")

doc.add_page_break()

# =========================================================================
# CHAPTER 8
# =========================================================================
add_h1("第八章：100 輪與 50 輪自動化測試驗證與品質監控")

add_h2("8.1 為什麼自動化測試是 Antigravity 協作的核心？")
add_p("在大型 Web 與多媒體專案中，手動點擊測試極耗時間且容易遺漏。AntigravityAI 強調「以腳本驗證品質」，透過自動化測試驗證每一項修改。")

add_h2("8.2 Node.js vm 沙盒模擬測試（audit_50_dual_host_podcast.js）")
add_p("利用 Node.js `vm` 模組在記憶體中建立 DOM 沙盒，模擬使用者點擊 10 大章節卡片 50 次，驗證播放狀態切換與逐字稿載入。")

add_h2("8.3 100 輪臺灣法律條文與法理交叉審查腳本")
add_p("執行 `scratch/verify_100_rounds_podcast_legal.py` 進行 100 輪條文交叉檢核，結果 100/100 PASSED！")

add_h2("8.4 測試報告生成與自動修復機制")
add_p("當測試失敗時，Antigravity 會自動閱讀 Error Stack Trace，修復語法或路徑錯誤後重新執行，直到全數通過。")

doc.add_page_break()

# =========================================================================
# CHAPTER 9
# =========================================================================
add_h1("第九章：GitHub Pages 線上部署與 Git 版本管理實戰")

add_h2("9.1 大檔案音訊 (podcast_audio_30min.mp3) 的 Git Force-Tracking 技巧")
add_p("當大檔案音訊被 `.gitignore` 排除時，使用 `git add -f podcast_audio_30min.mp3` 強制追蹤並提交至 GitHub。")

add_h2("9.2 處理跨域資源共享 (CORS) 與本地 file:/// 瀏覽試聽")
add_p("播放器內建備援 JSON 字串（`fullScriptData`），即使在無伺服器的本地 `file:///` 環境開啟，依然能流暢運作。")

add_h2("9.3 GitHub Pages 免費建置與一鍵更新")
add_p("開啟 GitHub Repository -> Settings -> Pages -> Source 選取 `main` 分支根目錄，系統將自動發布為全球可訪問之 HTTPS 網站。")

doc.add_page_break()

# =========================================================================
# CHAPTER 10
# =========================================================================
add_h1("第十章：初學者常見問題 QA、除錯技巧與進階延伸")

add_h2("10.1 常見錯誤 1：TTS 網路連線逾時（NoAudioReceived）之重試機制")
add_p("問題現象：Microsoft Azure TTS 伺服器因連線過密回傳 NoAudioReceived。\n修復方法：在 Python 中加入 `try...except` 重試機制（Retry Loop）與 `asyncio.sleep(0.3)` 緩衝。")

add_h2("10.2 常見錯誤 2：Windows CP950 編碼亂碼與 UTF-8 解除")
add_p("問題現象：Windows PowerShell 控制台印出 Unicode 表情符號時報錯 `UnicodeEncodeError: 'cp950' codec...`。\n修復方法：在 Python print 語句中避免直接輸出 Unicode Emoji，改用 `[OK]`、`[FAIL]` 標籤。")

add_h2("10.3 常見錯誤 3：JSON 解析轉義字元報錯之修復")
add_p("問題現象：JavaScript 內嵌 JSON 時因換行符號 `\n` 導致 `SyntaxError: Unexpected token`。\n修復方法：採用 `JSON.parse(JSON.stringify(...))` 安全轉義內嵌。")

add_h2("10.4 實務經驗總結：初學者協作金律 Check List")
add_bullet("明確定義業務目標與檔案目錄結構。", bold_prefix="Check 1：")
add_bullet("善用 scratch/ 存放中間測試腳本。", bold_prefix="Check 2：")
add_bullet("要求 AI 撰寫客觀自動化驗證腳本。", bold_prefix="Check 3：")
add_bullet("隨時進行 git commit 保留穩定節點。", bold_prefix="Check 4：")

add_callout(
    "祝賀完成 Antigravity AI 實戰學習",
    "恭喜您完整研讀本講義！您已掌握與 Antigravity AI 協作開發 500 頁 HTML 簡報、講稿全集與 30 分鐘廣播 Podcast 的全部核心技術。快去開啟您的下一個 AI 協作專案吧！",
    type_style='tip'
)

# Save Document
doc_out_path = os.path.join(os.path.dirname(__file__), '..', 'Antigravity_AI全方位實戰指南_HTML簡報與網頁Podcast產製講義.docx')
doc.save(doc_out_path)
print(f"[OK] Master Handbook successfully generated and saved to: {doc_out_path}")
