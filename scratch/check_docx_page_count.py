import os
import docx

doc_path = os.path.join(os.path.dirname(__file__), '..', 'Antigravity_AI全方位實戰指南_HTML簡報與網頁Podcast產製講義.docx')

doc = docx.Document(doc_path)
total_chars = 0
for p in doc.paragraphs:
    total_chars += len(p.text)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            total_chars += len(cell.text)

print(f"Total Paragraphs: {len(doc.paragraphs)}")
print(f"Total Tables: {len(doc.tables)}")
print(f"Total Characters: {total_chars}")
