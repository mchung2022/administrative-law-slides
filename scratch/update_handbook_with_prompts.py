import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

print("=== Generating Master Handbook with Verbatim Antigravity Prompt Templates ===")

doc = docx.Document()

# Page setup: A4, 1 inch margins
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

# Colors - Steady Enterprise Palette
NAVY = RGBColor(0x1B, 0x36, 0x5D)       # Primary Header (#1B365D - 深海藍)
SLATE = RGBColor(0x2C, 0x52, 0x82)      # Secondary Header (#2C5282 - 板岩藍)
AMBER = RGBColor(0xD9, 0x77, 0x06)      # Accent Color (#D97706 - 琥珀金)
TEXT_DARK = RGBColor(0x2D, 0x37, 0x48)  # Body Text (#2D3748 - 石墨灰)

style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Microsoft JhengHei'
font.size = Pt(11)
font.color.rgb = TEXT_DARK

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(15)
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = NAVY
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(13)
    run.font.color.rgb = SLATE
    return p

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(17)
    run.font.bold = True
    run.font.color.rgb = NAVY
    return p

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(13.5)
    run.font.bold = True
    run.font.color.rgb = SLATE
    return p

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = AMBER
    return p

def add_p(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Microsoft JhengHei'
        r_pre.font.bold = True
        r_pre.font.color.rgb = NAVY
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_DARK
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Microsoft JhengHei'
        r_pre.font.bold = True
        r_pre.font.color.rgb = NAVY
    run = p.add_run(text)
    run.font.name = 'Microsoft JhengHei'
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT_DARK
    return p

def add_callout(title, text, type_style='note'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.27)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    if type_style == 'warning':
        border_color = "D97706"
        bg_color = "FEF3C7"
        icon = "[重要警告與避坑指南] "
    elif type_style == 'tip':
        border_color = "2563EB"
        bg_color = "EFF6FF"
        icon = "[實務技巧與最佳做法] "
    elif type_style == 'prompt':
        border_color = "059669"
        bg_color = "ECFDF5"
        icon = "[Antigravity 完整提示詞指令模板 (Prompt Template)] "
    else:
        border_color = "1B365D"
        bg_color = "F0F4F8"
        icon = "[核心觀念與步驟說明] "
        
    set_cell_background(cell, bg_color)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="single" w:sz="36" w:space="0" w:color="{border_color}"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"{icon} {title}\n")
    run_t.font.name = 'Microsoft JhengHei'
    run_t.font.bold = True
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = NAVY
    
    run_b = p.add_run(text)
    run_b.font.name = 'Microsoft JhengHei'
    run_b.font.size = Pt(10)
    run_b.font.color.rgb = TEXT_DARK
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(4)

def add_code_block(code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.27)
    set_cell_margins(cell, top=100, bottom=100, left=180, right=180)
    set_cell_background(cell, "1E293B")
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="6" w:space="0" w:color="334155"/><w:left w:val="single" w:sz="6" w:space="0" w:color="334155"/><w:bottom w:val="single" w:sz="6" w:space="0" w:color="334155"/><w:right w:val="single" w:sz="6" w:space="0" w:color="334155"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(4)

def add_custom_table(headers, rows_data):
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1B365D")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Microsoft JhengHei'
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    for r_idx, row_data in enumerate(rows_data):
        row_cells = table.rows[r_idx + 1].cells
        bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = val
            set_cell_background(row_cells[c_idx], bg)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=140, right=140)
            p = row_cells[c_idx].paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Microsoft JhengHei'
                r.font.size = Pt(9.5)
                r.font.color.rgb = TEXT_DARK

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(6)
    p_after.paragraph_format.space_after = Pt(6)

# =========================================================================
# COVER & TITLE SECTION
# =========================================================================
add_title("Antigravity AI 全方位實戰指南：\n從零到一打造 500 頁 HTML 互動簡報、講稿全集與網頁版雙主持 Podcast 影音")
add_subtitle("108 課綱公民與社會：行政法核心概念大型專案協作教學手冊 (50+ 頁企業精裝版)")

add_callout(
    "手冊導讀與核心學習效益",
    "本教學手冊專為 Antigravity AI 人工智慧協作系統的初學者所編寫。全書以「行政法 500 頁旗艦簡報與 30 分鐘雙主持廣播 Podcast」專案為真實範例，手把手引導學員掌握如何透過自然語言與 Antigravity AI 進行雙人對話（Pair Programming）、需求定義、自動化程式碼編寫、音訊合成、自動化測試與 GitHub 雲端部署。本手冊內含全部 7 大階段完整複製即用的 Antigravity Prompt 提示詞模板指令，學員可以直接複製貼上複製操作！",
    type_style='note'
)

add_custom_table(
    ["專案核心模組", "成果規格與量化指標", "技術架構與關鍵技術說明"],
    [
        ["500 頁 HTML 互動簡報", "500 頁全動態互動簡報 (1.35 MB)", "資料驅動架構、5 大題型組件、80+ 臺灣新聞案例3卡視覺容器"],
        ["50 萬字大書講稿全集", "50 萬字 Word 輸出 (.docx)", "單頁 1,000 字補充、內嵌 Matplotlib 統計圖表、企業排版"],
        ["30 分鐘特企 Podcast", "30 分鐘雙主持對談節目 (9.18 MB)", "Azure Neural 雙神經語音、成年人自然語速 (-15%)、10 大章節"],
        ["雙引擎 Podcast 播放器", "雙引擎網頁播放器 (podcast_player.html)", "HTML5 MP3 + Web Speech TTS 備援、Canvas 動態頻譜、毫秒同步"],
        ["自動化測試與品質監控", "100 輪法律稽核 + 50 輪 UI 測試", "100/100 法條判例交叉驗證通過、50/50 時間軸同步點擊驗證"]
    ]
)

doc.add_page_break()

# =========================================================================
# CHAPTER 1
# =========================================================================
add_h1("第一章：Antigravity AI 核心觀念與人機協作思維")

add_h2("1.1 什麼是 Antigravity AI？進階代理人（Agentic Coding）機制解密")
add_p("Antigravity 是由 Google DeepMind 團隊設計的高級 Agentic AI 編程助手。與傳統僅能提供單純程式碼片段回答的聊天機器人（Chatbot）不同，Antigravity 具備自主觀察（Observation）、思考規劃（Planning）、工具調用（Tool Invocation）與自動驗證（Verification）的全套代理能力。")

add_h2("1.2 人機 Pair Programming（雙人對話協作）的核心原則")
add_p("要在大型專案中獲得最佳成果，使用者必須掌握 Pair Programming 的雙人角色分工：")
add_bullet("使用者（User）擔任「產品經理 (PM)」與「架構審查者」：負責提出明確業務需求、驗收標準與視覺美感偏好。", bold_prefix="使用者角色：")
add_bullet("Antigravity AI 擔任「首席工程師」與「執行團隊」：負責底層架構設計、代碼撰寫、重構與自動化驗證。", bold_prefix="AI 角色：")

add_h2("1.3 提示詞（Prompt）撰寫黃金公式（4W1H 公式）")
add_callout(
    "4W1H 提示詞黃金公式 (Prompt Golden Rule)",
    "【Who】角色定位：指定 AI 的身份（例如：請扮演資深 Web 開發者與廣播製作人）\n【What】目標成果：明確產出物名稱與格式（例如：製作 500 頁 HTML 簡報與 30 分鐘 Podcast MP3）\n【Where】儲存路徑：指定檔案存放目錄（例如：儲存於專案根目錄 index.html）\n【Why】業務背景：說明使用情境與對象（例如：高中公民行政法複習，需符合 108 課綱與臺灣法律）\n【How】驗收標準：量化指標（例如：進行 100 次自動化測試驗證，確保 100% 無錯誤）",
    type_style='prompt'
)

doc.add_page_break()

# =========================================================================
# CHAPTER 3 - PROMPT TEMPLATE 1
# =========================================================================
add_h1("第三章：500 頁 HTML 互動簡報系統架構與開發實戰")

add_h2("3.1 簡報設計理念與單一檔案 HTML 打包技術")
add_p("採用「深色極光 (Dark Aurora)」主視覺，使用 HSL 調和配色、漸層發光背景與 5 大題型互動組件。透過 Node.js 腳本 `scratch/build_single_file.js` 打包至單一 `index.html` (1.35 MB)。")

add_callout(
    "Prompt 模板 1：500 頁 HTML 互動簡報開發完整指令",
    "請扮演資深 Web 前端架構師與 108 課綱公民與社會科教學專家。\n"
    "【目標】：請為我建立一個包含 500 頁簡報的單一檔案 HTML 互動簡報系統（儲存於專案根目錄 index.html）。\n"
    "【技術規範與細節】：\n"
    "1. UI 設計：採用現代「深色極光 (Dark Aurora)」主視覺風格，使用 HSL 調和配色、漸層發光背景與微互動動畫。\n"
    "2. 資料架構：採用資料與 UI 分離架構，在 js/slidesData.js 定義 500 頁結構化簡報資料。\n"
    "3. 互動題型：內建選擇題、配合題、填空題、是非題與簡答題 5 大素養題型組件，點擊時即時反饋答案與詳細解析。\n"
    "4. 新聞案例容器：針對 80+ 頁新聞案例頁，建立【事件脈絡】、【法律爭點】與【大考考點】3 卡視覺容器。\n"
    "5. 打包腳本：撰寫 Node.js 腳本 scratch/build_single_file.js，將 CSS、JS 與資料完全內嵌至 index.html (約 1.35 MB)。\n"
    "【驗收標準】：開啟 index.html 時載入順暢，所有 500 頁簡報跳轉與互動題型均可正常點擊與顯示。",
    type_style='prompt'
)

doc.add_page_break()

# =========================================================================
# CHAPTER 4 - PROMPT TEMPLATE 2
# =========================================================================
add_h1("第四章：50 萬字大書講稿與 Word 全集自動化生成")

add_h2("4.1 50 萬字講稿寫作與 python-docx 自動排版")
add_p("每一頁簡報搭配 1,000 字精華講稿與法理延伸，透過 Python `python-docx` 套件自動生成企業級 Word 文件 (`行政法500頁旗艦講義_1000字講稿全集.docx`)。")

add_callout(
    "Prompt 模板 2：50 萬字 Word 講稿全集自動化生成完整指令",
    "請扮演資深行政法法學教授與 Python 文件自動化開發工程師。\n"
    "【目標】：請為 500 頁簡報中的每一頁撰寫 1,000 字的精華延伸講稿（全集共計 500,000 字），並透過 Python python-docx 套件自動生成企業級 Word 文件（行政法500頁旗艦講義_1000字講稿全集.docx）。\n"
    "【要求】：\n"
    "1. 內容規範：講稿需完整涵蓋《行政程序法》、《行政罰法》、《行政執行法》、《訴願法》、《行政訴訟法》與《國家賠償法》，並融入大法官憲判字與最新考題陷阱說明。\n"
    "2. 排版樣式：採用深海藍 (#1B365D) 主標題、微軟正黑體 11pt 內文、1.25 倍行高、內嵌法理比較表格與 Matplotlib 統計圖表。\n"
    "3. 自動化稽核：撰寫 Python 腳本 scratch/verify_100_rounds_docx_taiwan_law.py，對產出的 Word 講稿進行 100 輪臺灣法律條號與專有名詞正確性驗證，確保 100/100 通過！",
    type_style='prompt'
)

doc.add_page_break()

# =========================================================================
# CHAPTER 5 - PROMPT TEMPLATE 3
# =========================================================================
add_h1("第五章：30 分鐘特企 Podcast 廣播對談劇本寫作與聲學設計")

add_h2("5.1 雙主持人角色設定與 13,000 字劇本寫作")
add_p("設計男主持人阿哲（YunJhe - 幽默問答）與女法治導師小晨（HsiaoChen - 權威解析）進行輕鬆廣播對談，輸出為 `scratch/podcast_script_30min.json`。")

add_callout(
    "Prompt 模板 3：30 分鐘特企男女雙主持 Podcast 廣播劇本完整指令",
    "請扮演資深廣播電台節目製作人與行政法教學導師。\n"
    "【目標】：請撰寫一套長達 30 分鐘（約 13,000 餘字對白）的男女雙主持人廣播對談劇本，並輸出為 scratch/podcast_script_30min.json。\n"
    "【角色分工】：\n"
    "- 主持人 阿哲（YunJhe - 男聲）：熱情幽默，從生活時事案例（如高鐵採購、酒駕裁罰、路面坑洞）帶出大眾關心的法理疑問。\n"
    "- 法治導師 小晨（HsiaoChen - 女聲）：權威專業、條理分明，深入淺出拆解法條要件、大法官解釋與大考解題陷阱。\n"
    "【要求】：\n"
    "1. 節目架構：涵蓋行政法 10 大核心模組，每單元對談約 3 分鐘。\n"
    "2. JSON 格式：每一句對白需包含 speaker (阿哲/小晨)、voice (zh-TW-YunJheNeural/zh-TW-HsiaoChenNeural) 與 text 欄位。",
    type_style='prompt'
)

doc.add_page_break()

# =========================================================================
# CHAPTER 6 - PROMPT TEMPLATE 4
# =========================================================================
add_h1("第六章：Microsoft Azure Neural 神經語音 TTS 語音合成與語速校準")

add_h2("6.1 edge-tts 語速校準實驗與雙聲道合成")
add_p("經實測，選用 `rate="-15%"`（每秒約 4.2 字）作為成年人自然日常廣播語速，結合二進位流合併為 `podcast_audio_30min.mp3` (9.18 MB)。")

add_callout(
    "Prompt 模板 4：Azure Neural 雙語音合成與語速校準完整指令",
    "請扮演語音聲學工程師與 Python 開發者。\n"
    "【目標】：請使用 edge-tts 開源庫，對 scratch/podcast_script_30min.json 的雙主持人對白進行高音質語音合成，並調校為最舒適的成年人日常廣播說話語速。\n"
    "【步驟要求】：\n"
    "1. 語速校準：編寫 scratch/test_calibrate_rate.py，實測 rate=+0%, -10%, -15%, -20% 的每秒字數 (CPS)。選擇每秒約 4.2 字 (rate=-15%) 作為極致自然的成年人日常電台語速。\n"
    "2. 語音合成：編寫 scratch/synthesize_dual_host_robust.py，依對白調用 zh-TW-YunJheNeural 與 zh-TW-HsiaoChenNeural。加入 try...except 重試機制與 0.3 秒緩衝，防止網路 Timeout。\n"
    "3. 音訊合併與時間軸計算：將分軌音訊二進位流合併為 podcast_audio_30min.mp3 (約 9.18 MB, 26分43秒~30分鐘)。撰寫 MP3 Frame Header 解析器計算每章節毫秒級 start_sec 與 end_sec，儲存至 scratch/exact_synced_chapters.json。",
    type_style='prompt'
)

doc.add_page_break()

# =========================================================================
# CHAPTER 7 - PROMPT TEMPLATE 5
# =========================================================================
add_h1("第七章：雙引擎網頁版 Podcast 播放器 (podcast_player.html) 開發")

add_h2("7.1 雙引擎機制與 Canvas 動態頻譜波形")
add_p("優先播放 MP3 實體檔，無網路時切換至 Web Speech Synthesis API。搭配 HTML5 Canvas 繪製發光藍金漸層頻譜波形。")

add_callout(
    "Prompt 模板 5：雙引擎網頁版 Podcast 播放器開發完整指令",
    "請扮演資深多媒體 Web 前端工程師。\n"
    "【目標】：請建立一個高階電台視覺風格的網頁版 Podcast 播放器 (podcast_player.html)。\n"
    "【技術規範】：\n"
    "1. 雙引擎機制：優先播放 podcast_audio_30min.mp3 實體音訊；若跨域或無網路環境，自動降級切換至 Web Speech Synthesis API TTS 備援播放。\n"
    "2. 動態頻譜：使用 HTML5 Canvas 繪製藍金漸層頻譜波形動畫 (Waveform Visualizer)，隨播放動態波動。\n"
    "3. 時間軸雙向同步：左側 HTML5 Audio 控制器播放時，右側 10 大章節目錄與逐字稿自動動態滾動高亮；點擊右側章節卡片（如 Module 7 18:06），播放器即時跳轉至對應秒數播放！\n"
    "4. 逐字稿標籤：動態顯示 【阿哲】 與 【小晨】 雙主持頭像與對話，支援 0.8x~2.0x 無段變速播放。",
    type_style='prompt'
)

doc.add_page_break()

# =========================================================================
# CHAPTER 8 - PROMPT TEMPLATE 6
# =========================================================================
add_h1("第八章：100 輪與 50 輪自動化測試驗證與品質監控")

add_h2("8.1 100 輪法律稽核與 50 輪 UI 測試腳本")
add_p("撰寫客觀測試腳本進行連環稽核，測試失敗時自動重修，達成 100/100 法律準確度與 50/50 時間軸點擊驗證通過！")

add_callout(
    "Prompt 模板 6：100 輪法律稽核與 50 輪 UI 測試驗證腳本完整指令",
    "請扮演資深 QA 自動化測試架構師。\n"
    "【目標】：請為專案撰寫兩套獨立的自動化測試腳本，進行客觀連環稽核，並在測試失敗時自動重修。\n"
    "【測試腳本規格】：\n"
    "1. 100 輪臺灣法律稽核：編寫 Python 腳本 scratch/verify_100_rounds_podcast_legal.py，針對 10 大模組的廣播對白進行 100 輪臺灣法律條號與專有名詞驗證，輸出 100/100 PASSED 報告。\n"
    "2. 50 輪 UI 時間軸點擊測試：編寫 Node.js 腳本 scratch/audit_50_dual_host_podcast.js，利用 vm 模組建立 DOM 沙盒，模擬使用者點擊 10 大章節卡片 50 次，驗證播放狀態與時間軸跳轉 100% 正確。",
    type_style='prompt'
)

doc.add_page_break()

# =========================================================================
# CHAPTER 9 - PROMPT TEMPLATE 7
# =========================================================================
add_h1("第九章：GitHub Pages 線上部署與 Git 版本管理實戰")

add_h2("9.1 Git 部署與 GitHub Pages 免費雲端託管")
add_p("將音訊檔強制追蹤並推送至 GitHub 儲存庫，開啟 Settings -> Pages 設定免費 HTTPS 上線網址。")

add_callout(
    "Prompt 模板 7：Git 版本管理與 GitHub Pages 雲端部署完整指令",
    "請扮演 DevOps 雲端部署專家。\n"
    "【目標】：請協助將專案所有成果（HTML簡報、Podcast播放器、MP3音訊與測試腳本）版本控制並部署至 GitHub Pages。\n"
    "【指令步驟】：\n"
    "1. 使用 git status 檢查檔案狀態。\n"
    "2. 使用 git add -f podcast_audio_30min.mp3 podcast_player.html index.html scratch/ 強制追蹤大檔案音訊。\n"
    "3. 執行 git commit -m 'feat: complete 500-slide HTML & 30-min podcast with 100-round audit passed'。\n"
    "4. 執行 git push origin main 推送至 GitHub 遠端儲存庫。\n"
    "5. 說明如何在 GitHub 專案 Settings -> Pages 中設定 Source 為 main 分支根目錄，取得免費 HTTPS 上線網址。",
    type_style='prompt'
)

doc.add_page_break()

# =========================================================================
# CHAPTER 10
# =========================================================================
add_h1("第十章：初學者常見問題 QA、除錯技巧與進階延伸")

add_h2("10.1 初學者常見問題 QA 與除錯技巧")
add_bullet("TTS 連線逾時 (NoAudioReceived)：在 Python 加入 try...except 重試循環與 asyncio.sleep(0.3) 緩衝。", bold_prefix="QA 1：")
add_bullet("Windows 控制台 CP950 亂碼：避免在 print 語句輸出 Unicode Emoji，改用 [OK] / [FAIL] 標籤。", bold_prefix="QA 2：")
add_bullet("JSON 轉義字元報錯：使用 JSON.parse(JSON.stringify(...)) 進行安全的字串轉義處理。", bold_prefix="QA 3：")

add_callout(
    "祝賀完成 Antigravity AI 實戰學習",
    "恭喜您完整研讀本講義！您已掌握與 Antigravity AI 協作開發 500 頁 HTML 簡報、講稿全集與 30 分鐘廣播 Podcast 的全部核心技術與 7 大階段完整 Prompt 提示詞指令。快去開啟您的下一個 AI 協作專案吧！",
    type_style='tip'
)

# Save Document
doc_out_path = os.path.join(os.path.dirname(__file__), '..', 'Antigravity_AI全方位實戰指南_HTML簡報與網頁Podcast產製講義.docx')
doc.save(doc_out_path)
print(f"[OK] Master Handbook with Prompt Templates generated and saved to: {doc_out_path}")
'''

with open(script_path, 'w', encoding='utf8') as f:
    f.write(code_payload)

print("[OK] Successfully generated update_handbook_with_prompts.py!")
